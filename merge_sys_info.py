import time
from datetime import datetime as dt, timedelta
from platform import node
from struct import unpack
from winreg import OpenKeyEx, QueryValueEx, HKEY_LOCAL_MACHINE, QueryInfoKey, EnumKey, KEY_READ
from subprocess import check_output
from wmi import WMI

from docx import Document
from windows_tools import product_key


def get_size(bts: int, ending='iB') -> str:
    size = 1024
    for item in ["", "K", "M", "G", "T", "P"]:
        if bts < size:
            return f"{bts:.2f} {item}{ending}" if bts > 0 else f"{bts:.2f} {item}B"
        bts /= size



def winreg_os() -> (dict):
    try:
        win_info = dict()
        if comp_info := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\ComputerName\ComputerName"):
            win_info.update({'ComputerName': QueryValueEx(comp_info, 'ComputerName')[0]})
        if comp_shutdown := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\Windows"):
            shutdown_time_bin = QueryValueEx(comp_shutdown, 'ShutdownTime')[0]
            shutdown_time = (dt(1601, 1, 1) + timedelta(microseconds=float(unpack("<Q", shutdown_time_bin)[0]) / 10)). \
                strftime('%Y-%m-%d %H:%M:%S')
            win_info.update({'ShutdownTime': shutdown_time})
        if win_ver := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion"):
            for key in ["ProductName", "EditionID", "DisplayVersion", "CurrentBuild", "UBR", "InstallDate",
                        "RegisteredOwner"]:
                try:
                    if key == "InstallDate":
                        win_info.update({key: str(dt.fromtimestamp(QueryValueEx(win_ver, f'{key}')[0]))})
                    else:
                        win_info.update({key: QueryValueEx(win_ver, f'{key}')[0]})
                except FileNotFoundError:
                    continue
        if tz_key := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\TimeZoneInformation"):
            win_info.update({"TimeZone": QueryValueEx(tz_key, 'TimeZoneKeyName')[0]})
        if pkey := product_key.get_windows_product_key_from_reg():
            win_info.update({"ActivateKey": pkey})
        elif pkey := product_key.get_windows_product_key_from_wmi():
            win_info.update({"ActivateKey": pkey})
        else:
            win_info.update({"ActivateKey": "No Key"})
        return win_info if win_info else False
    except Exception:
        return False



def wmic_os() -> (dict):
    os_info = dict()
    try:
        os_i = check_output("wmic OS get Caption, InstallDate, Version, WindowsDirectory, LastBootUpTime /value",
                            shell=True).decode().strip()
    except UnicodeDecodeError:
        os_i = check_output("wmic OS get Caption, InstallDate, Version, WindowsDirectory, LastBootUpTime /value",
                            shell=True).decode("cp866").strip()
    if os_i:
        out = [{x.strip().split("=")[0]: x.strip().split("=")[1]} for x in os_i.splitlines() if x.strip()]
        if comp_name := check_output("wmic computersystem get name /value", shell=True).decode().strip():
            os_info.update({"ComputerName": comp_name.split("=")[1].strip()})
        for item in out:
            for it in item:
                if item[it]:
                    if it == "InstallDate":
                        date = dt.strptime(item[it][:len(item[it]) - 4], '%Y%m%d%H%M%S.%f').strftime(
                            "%Y-%m-%d %H:%M:%S")
                        item = {"InstallDate": date}
                    if it == "LastBootUpTime":
                        date = dt.strptime(item[it][:len(item[it]) - 4], '%Y%m%d%H%M%S.%f').strftime(
                            "%Y-%m-%d %H:%M:%S")
                        item = {"LastBootUpTime": date}
                    os_info.update(item)
    if tz_get := check_output("wmic TIMEZONE get Caption /value", shell=True).decode().strip():
        tz = [{"TimeZone": x.strip().split("=")[1]} for x in tz_get.splitlines() if x.strip()]
        os_info.update(tz[0])
    if user_get := check_output('wmic USERACCOUNT Where (Status="OK") get Name /value', shell=True).decode().strip():
        user = [{"UserName": x.strip().split("=")[1]} for x in user_get.splitlines() if x.strip()]
        os_info.update(user[0])
    return os_info if os_info else False



def bios_winreg() -> (dict):
    md_dict = dict()
    if sbv := OpenKeyEx(HKEY_LOCAL_MACHINE, r"HARDWARE\DESCRIPTION\System"):
        md_dict.update({"SystemBiosVersion": QueryValueEx(sbv, "SystemBiosVersion")[0][0]})
    for key in ["BIOSVendor", "BIOSVersion", "BIOSReleaseDate"]:
        if bios := OpenKeyEx(HKEY_LOCAL_MACHINE, r"HARDWARE\DESCRIPTION\System\BIOS"):
            try:
                md_dict.update({key: QueryValueEx(bios, key)[0]})
            except FileNotFoundError:
                continue
        else:
            return False
    return md_dict if md_dict else False



def bios_wmic() -> (dict):
    bios_info = dict()
    if bios := check_output("wmic BIOS get Version, Manufacturer, Name /value", shell=True).decode(). \
            strip():
        out = [{f'{x.strip().split("=")[0]}': x.strip().split("=")[1]} for x in bios.splitlines() if x.strip()]
        for item in out:
            bios_info.update(item)
        return bios_info if bios_info else False
    return False



def motherboard_winreg() -> (dict):
    md_dict = dict()
    if mb_info := OpenKeyEx(HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\SystemInformation"):
        md_dict.update({'SystemManufacturer': QueryValueEx(mb_info, 'SystemManufacturer')[0]})
        md_dict.update({'SystemProductName': QueryValueEx(mb_info, 'SystemProductName')[0]})
        return md_dict if md_dict else False
    return False


def motherboard_wmic() -> (dict):
    mb_info = dict()
    if mb_get := check_output("wmic baseboard get Product, Manufacturer /value", shell=True).decode().strip():
        out = [{f'{x.strip().split("=")[0]}': x.strip().split("=")[1]} for x in mb_get.splitlines() if x.strip()]
        for item in out:
            mb_info.update(item)
        return mb_info if mb_info else False
    return False



def memory_wmic() -> (dict):
    memory_info = dict()
    if total_memory_get := check_output("wmic COMPUTERSYSTEM get TotalPhysicalMemory /value",
                                        shell=True).decode().strip():
        total_memory = [{x.strip().split("=")[0]: get_size(int(x.strip().split("=")[1]))}
                        for x in total_memory_get.splitlines() if x.strip()]
        memory_info.update(total_memory[0])
    if memory_tag_get := check_output("wmic MEMORYCHIP get Tag /value", shell=True).decode().strip():
        memory_tag = [x.strip().split("=")[1] for x in memory_tag_get.splitlines() if x.strip()]
        for tag in memory_tag:
            info_get = check_output(f'wmic MEMORYCHIP Where (Tag="{tag}") get Capacity, ConfiguredClockSpeed, '
                                    f'Manufacturer, PartNumber, SerialNumber /value', shell=True).decode().strip()
            memory_info[tag.strip()] = dict()
            for x in info_get.splitlines():
                if x.strip():
                    if x.strip().split("=")[0] == "Capacity" and x.strip().split("=")[1]:
                        memory_info[tag.strip()].update({
                            x.strip().split("=")[0].strip(): get_size(int(x.strip().split("=")[1])).strip()
                        })
                    else:
                        memory_info[tag.strip()].update({
                            x.strip().split("=")[0].strip(): x.strip().split("=")[1].strip()
                        })
    return memory_info if memory_info else False



def memory_wmi() -> (dict):
    memory_dict = dict()
    wmi_obj = WMI()
    if memory_data := wmi_obj.Win32_PhysicalMemory():
        memory_dict.update({"TotalPhysicalMemory": get_size(sum([int(mem.Capacity) for mem in memory_data]))})

        for i, mem in enumerate(memory_data):
            memory_dict[f"Physical Memory {i}"] = dict()
            memory_dict[f"Physical Memory {i}"].update({
                "Capacity": get_size(int(mem.Capacity)),
                "ConfiguredClockSpeed": mem.Speed,
                "Manufacturer": mem.Manufacturer,
                "PartNumber": mem.PartNumber,
                "SerialNumber": mem.SerialNumber
            })
    return memory_dict if memory_dict else False



def cpu_winreg():
    try:
        proc_info = dict()
        loc = "HARDWARE\\DESCRIPTION\\System\\CentralProcessor"
        with OpenKeyEx(HKEY_LOCAL_MACHINE, loc) as h_apps:
            if QueryInfoKey(h_apps)[0]:
                proc_info.update({"CoreCount": QueryInfoKey(h_apps)[0]})
                try:
                    core = OpenKeyEx(h_apps, EnumKey(h_apps, 0))
                    proc_info.update({
                        "ProcessorNameString": QueryValueEx(core, 'ProcessorNameString')[0].strip(),
                        "Identifier": QueryValueEx(core, 'Identifier')[0].strip(),
                        "VendorIdentifier": QueryValueEx(core, 'VendorIdentifier')[0].strip(),
                        "~MHz": QueryValueEx(core, '~MHz')[0]
                    })
                except FileNotFoundError:
                    return False
        return proc_info if proc_info else False
    except Exception:
        return False


def cpu_wmic() -> (dict):
    mb_info = dict()
    if mb_get := check_output("wmic cpu get Name, Caption, Manufacturer, SocketDesignation, MaxClockSpeed /value",
                              shell=True).decode().strip():
        out = [{f'{x.strip().split("=")[0]}': x.strip().split("=")[1]} for x in mb_get.splitlines() if x.strip()]
        for item in out:
            mb_info.update(item)
    if core_all := check_output("wmic COMPUTERSYSTEM get NumberOfLogicalProcessors /value",
                                shell=True).decode().strip():
        mb_info.update({'NumberOfPhysicalProcessors': str(int(int(core_all.split("=")[1].strip()) / 2)),
                        f'{core_all.split("=")[0].strip()}': core_all.split("=")[1].strip()})
    return mb_info if mb_info else False



def gpu_wmic() -> (dict):
    gpu = dict()
    if output := check_output('wmic path win32_VideoController get Name, AdapterRAM, VideoProcessor, '
                              'CurrentHorizontalResolution, CurrentVerticalResolution, CurrentRefreshRate /value',
                              shell=True).decode().strip():
        out = [{x.strip().split("=")[0]: x.strip().split("=")[1]} for x in output.splitlines() if x.strip()]
        res = ""
        for vrm in out:
            for vr in vrm:
                if vr == 'AdapterRAM':
                    gpu.update({"AdapterRAM": get_size(int(vrm[vr]))})
                    continue
                if vr == 'CurrentHorizontalResolution':
                    res += f'{vrm[vr]}x'
                    continue
                if vr == 'CurrentVerticalResolution':
                    res += str(vrm[vr])
                    gpu.update({"Resolution": res})
                    continue
                gpu.update(vrm)
        return gpu if gpu else False
    return False


def gpu_wmi() -> (dict):
    gpu = dict()
    if vc := WMI().Win32_VideoController()[0]:
        gpu.update({
            "Name": vc.Description,
            "AdapterRAM": get_size(abs(vc.AdapterRAM)),
            "Resolution": f'{vc.CurrentHorizontalResolution}x{vc.CurrentVerticalResolution}',
            "CurrentRefreshRate": f'{vc.CurrentRefreshRate} Гц',
            "VideoProcessor": vc.VideoProcessor
        })
    return gpu if gpu else False



def hdd_ssd_wmic() -> (dict):
    disk_info = dict()
    if caption_get := check_output("wmic diskdrive get Caption /value", shell=True).decode().strip():
        caption = [x.strip().split("=")[1] for x in caption_get.splitlines() if x.strip()]
        for num, cap in enumerate(caption):
            if info_get := check_output(
                    f'wmic diskdrive Where (Caption="{cap}") get DeviceID, FirmwareRevision, MediaType, '
                    f'Partitions, SerialNumber, Size /value', shell=True).decode().strip():
                disk_info[f'{cap} {num}'] = dict()
                disk_info[f'{cap} {num}'].update({"Product": cap})
                for x in info_get.splitlines():
                    if x.strip():
                        if x.strip().split("=")[0] == "Size" and x.strip().split("=")[1]:
                            disk_info[f'{cap} {num}'].update({
                                f'{x.strip().split("=")[0].strip()}': get_size(int(x.strip().split("=")[1].strip()))
                            })
                        else:
                            disk_info[f'{cap} {num}'].update({
                                f'{x.strip().split("=")[0].strip()}': x.strip().split("=")[1].strip()
                            })
    return disk_info if disk_info else False


def hdd_ssd_wmi() -> (dict):
    disk_info = dict()
    c = WMI()
    if disks := c.Win32_DiskDrive():
        for disk in disks:
            disk_info[disk.DeviceID] = {
                'Caption': disk.Model,
                'MediaType': disk.InterfaceType,
                'Capacity': get_size(int(disk.Size))
            }
    return disk_info if disk_info else False



def cdrom_wmic() -> (dict):
    cdrom_info = dict()
    if caption_get := check_output("wmic CDROM get Caption /value", shell=True).decode().strip():
        caption = [x.strip().split("=")[1] for x in caption_get.splitlines() if x.strip()]
        for num, cap in enumerate(caption):
            if info_get := check_output(
                    f'wmic CDROM Where (Caption="{cap}") get Drive, VolumeName, VolumeSerialNumber, Size '
                    f'/value', shell=True).decode().strip():
                cdrom_info[cap] = dict()
                cdrom_info[cap].update({"Product": cap})
                for x in info_get.splitlines():
                    if x.strip():
                        if x.strip().split("=")[0].strip() == "Size" and x.strip().split("=")[1].strip():
                            cdrom_info[cap].update({
                                f'{x.strip().split("=")[0].strip()}': get_size(int(x.strip().split("=")[1].strip()))
                            })
                        cdrom_info[cap].update({
                            f'{x.strip().split("=")[0].strip()}': x.strip().split("=")[1].strip()
                        })
    return cdrom_info if cdrom_info else False



def cdrom_wmi() -> (dict):
    cdrom_info = dict()
    c = WMI()
    if cdroms := c.Win32_CDROMDrive():
        for cdrom in cdroms:
            cdrom_info[cdrom.Caption] = dict()
            cdrom_info[cdrom.Caption].update({
                'Drive': cdrom.Caption,
                'MediaType': cdrom.MediaType,
                'Status': cdrom.Status,
                'SerialNumber': cdrom.SerialNumber,
                'Manufacturer': cdrom.Manufacturer
            })
    return cdrom_info if cdrom_info else False



def nic_winreg():
    try:
        nic = dict()
        loc = r'SOFTWARE\Microsoft\Windows NT\CurrentVersion\NetworkCards'
        desc = []
        if adapt_name := OpenKeyEx(HKEY_LOCAL_MACHINE, loc):
            for idx in range(QueryInfoKey(adapt_name)[0]):
                adapter = OpenKeyEx(adapt_name, EnumKey(adapt_name, idx))
                nic[QueryValueEx(adapter, 'ServiceName')[0]] = dict()
                nic[QueryValueEx(adapter, 'ServiceName')[0]].update({
                    "Description": QueryValueEx(adapter, 'Description')[0]
                })
                desc.append(QueryValueEx(adapter, 'Description')[0])
                loc_adapt = r'SYSTEM\ControlSet001\Control\Class\{4d36e972-e325-11ce-bfc1-08002be10318}'
                cfg = []
                if adapt := OpenKeyEx(HKEY_LOCAL_MACHINE, loc_adapt):
                    for idc in range(QueryInfoKey(adapt)[0]):
                        try:
                            adpt = OpenKeyEx(adapt, EnumKey(adapt, idc))
                            if QueryValueEx(adpt, 'DriverDesc')[0] in desc:
                                nic[QueryValueEx(adpt, 'NetCfgInstanceId')[0]].update({
                                    "Description": QueryValueEx(adpt, 'DriverDesc')[0]
                                })
                                cfg.append(QueryValueEx(adpt, 'NetCfgInstanceId')[0])
                        except (FileNotFoundError, PermissionError):
                            continue
                inter = r'SYSTEM\CurrentControlSet\Services\Tcpip\Parameters\Interfaces'
                if inter_cfg := OpenKeyEx(HKEY_LOCAL_MACHINE, inter):
                    for idb in range(QueryInfoKey(inter_cfg)[0]):
                        if EnumKey(inter_cfg, idb).upper() in cfg:
                            nic[EnumKey(inter_cfg, idb).upper()].update({
                                "NetCfgInstanceId": EnumKey(inter_cfg, idb).upper()})
                            intr = OpenKeyEx(inter_cfg, EnumKey(inter_cfg, idb))
                            try:
                                nic[EnumKey(inter_cfg, idb).upper()].update({
                                    "DhcpDefaultGateway": QueryValueEx(intr, 'DhcpDefaultGateway')[0]})
                            except FileNotFoundError:
                                pass
                            try:
                                nic[EnumKey(inter_cfg, idb).upper()].update({
                                    "DhcpIPAddress": QueryValueEx(intr, 'DhcpIPAddress')[0]})
                            except FileNotFoundError:
                                pass
                            try:
                                nic[EnumKey(inter_cfg, idb).upper()].update({
                                    "DhcpIPAddress": QueryValueEx(intr, 'DhcpIPAddress')[0]})
                            except FileNotFoundError:
                                pass
                            netw = r'SYSTEM\ControlSet001\Control\Network\{4D36E972-E325-11CE-BFC1-08002BE10318}' + '\\' + \
                                   EnumKey(inter_cfg, idb).upper() + '\\' + 'Connection'
                            if netw_cfg := OpenKeyEx(HKEY_LOCAL_MACHINE, netw):
                                nic[EnumKey(inter_cfg, idb).upper()].update({
                                    "Name": QueryValueEx(netw_cfg, 'Name')[0]})
        return nic if nic else False
    except Exception:
        return False


def nic_wmic() -> (dict):
    nic_info = dict()
    if description_get := check_output('wmic NIC Where (PhysicalAdapter="TRUE") get Description /value', shell=True). \
            decode().strip():
        description = [x.strip().split("=")[1] for x in description_get.splitlines() if x.strip()]
        for nic in description:
            try:
                nic_d = check_output(
                    f'wmic NIC Where (Description="{nic}") get MACAddress, Manufacturer, NetConnectionID'
                    f' /value', shell=True).decode().strip()
            except UnicodeDecodeError:
                nic_d = check_output(
                    f'wmic NIC Where (Description="{nic}") get MACAddress, Manufacturer, NetConnectionID'
                    f' /value', shell=True).decode("cp866").strip()
            try:
                nicconfig_d = check_output(
                    f'wmic NICCONFIG Where (Description="{nic}") get DHCPServer, IPAddress /value',
                    shell=True).decode().strip()
            except UnicodeDecodeError:
                nicconfig_d = check_output(
                    f'wmic NICCONFIG Where (Description="{nic}") get DHCPServer, IPAddress /value',
                    shell=True).decode("cp866").strip()
            nic_info[nic.strip()] = dict()
            nic_info[nic.strip()].update({"Product": nic.strip()})
            if nic_d:
                for x in nic_d.splitlines():
                    if x.strip():
                        nic_info[nic.strip()].update({
                            x.strip().split("=")[0].strip(): x.strip().split("=")[1].strip()
                        })
            if nicconfig_d:
                for x in nicconfig_d.splitlines():
                    if x.strip():
                        nic_info[nic.strip()].update({
                            x.strip().split("=")[0].strip(): x.strip().split("=")[1].strip()
                        })
    return nic_info if nic_info else False


wmic_info = ""


def print_wmic(part, dict_info):
    global wmic_info
    synonyms = {"ComputerName": "Имя компьютера", "Caption": "Название", "InstallDate": "Дата установки",
                "LastBootUpTime": "Время последней загрузки", "Version": "Версия",
                "WindowsDirectory": "Директория Windows", "TimeZone": "Часовой пояс", "UserName": "Имя пользователя",
                "Manufacturer": "Производитель", "Name": "Название", "Product": "Изделие",
                "MaxClockSpeed": "Максимальная тактовая частота", "SocketDesignation": "Название сокета",
                "NumberOfPhysicalProcessors": "Количество физических процессоров", "VideoProcessor": "Видеопроцессор",
                "NumberOfLogicalProcessors": "Количество логических процессоров", "Capacity": "Емкость",
                "AdapterRAM": "Оперативная память адаптера", "CurrentRefreshRate": "Текущая частота обновления",
                "Resolution": "Разрешение", "TotalPhysicalMemory": "Общий объем физической памяти", "Socket": "Сокет",
                "ConfiguredClockSpeed": "Настроенная тактовая частота", "PartNumber": "Номер партии",
                "SerialNumber": "Серийный номер", "DeviceID": "Идентификатор устройства", "MediaType": "Тип носителя",
                "FirmwareRevision": "Ревизия прошивки", "Partitions": "Разделы", "Size": "Объем", "Drive": "Диск",
                "VolumeName": "Имя тома", "VolumeSerialNumber": "Серийный номер тома", "MACAddress": "MAC-адрес",
                "NetConnectionID": "Идентификатор сетевого подключения", "DHCPServer": "DHCP-сервер",
                "IPAddress": "IP-адрес", "BuildNumber": "Номер сборки", "ID": "Идентификатор", "Status": "Статус",
                "DefaultIPGateway": "IP-адрес шлюза по-умолчанию", "DNSHostName": "DNS Имя хоста",
                "IPv4Address": "IPv4-адрес", "IPv6Address": "IPv6-адрес", "IPSubnet": "Маска подсети",
                "ServiceName": "Название службы", "CurrentBuild": "Текущая сборка", "UBR": "Номер версии",
                "RegisteredOwner": "Имя пользователя", "ActivateKey": "Ключ активации",
                "SystemBiosVersion": "Версия Bios системы", "BIOSVendor": "Производитель", "BIOSVersion": "Версия",
                "BIOSReleaseDate": "Дата выпуска релиза", "ShutdownTime": "Время выключения", "ProductName": "Название",
                "EditionID": "Идентификатор редакции", "DisplayVersion": "Версия для отображения",
                "SystemManufacturer": "Производитель", "SystemProductName": "Название сокета",
                "CoreCount": "Количество ядер", "ProcessorNameString": "Название", "Identifier": "Идентификатор",
                "VendorIdentifier": "Производитель", "~MHz": "Тактовая частота", "Vendor": "Производитель",
                "Model": "Модель", "Revision": "Ревизия", "Description": "Название",
                "NetCfgInstanceId": "Идентификатор", "DhcpDefaultGateway": "Шлюз по-умолчанию",
                "DhcpIPAddress": "IP-адрес"}
    part += f'{"-" * 50}\n'
    for key in dict_info:
        if type(dict_info[key]) == dict:
            for item in dict_info[key]:
                part += f'{synonyms[item]}: {dict_info[key][item]}\n'
            part += "\n"
        else:
            part += f'{synonyms[key]}: {dict_info[key]}\n'
    print(part)
    wmic_info += f'{part}\n'


def main():
    global wmic_info
    t = time.monotonic()
    document = Document()
    document.add_heading(f'Сводная информация о компьютере: {node()}')

    if os_info := winreg_os():
        print_wmic("Информация об операционной системе\n", os_info)
    else:
        if os_wmic := wmic_os():
            print_wmic("Информация об операционной системе\n", os_wmic)
    if bios_info := bios_winreg():
        print_wmic("Информация о BIOS\n", bios_info)
    else:
        if wmic_bios := bios_wmic():
            print_wmic("Информация о BIOS\n", wmic_bios)
    if mb_info := motherboard_winreg():
        print_wmic("Информация о материнской плате\n", mb_info)
    else:
        if wmic_mb := motherboard_wmic():
            print_wmic("Информация о материнской плате\n", wmic_mb)
    if memory_info := memory_wmic():
        print_wmic("Информация об оперативной памяти\n", memory_info)
    else:
        if mem_wmi := memory_wmi():
            print_wmic("Информация об оперативной памяти\n", mem_wmi)
    if cpu_info := cpu_winreg():
        print_wmic("Информация о процессоре\n", cpu_info)
    else:
        if wmic_cpu := cpu_wmic():
            print_wmic("Информация о процессоре\n", wmic_cpu)
    if gpu_info := gpu_wmic():
        print_wmic("Информация о видеокарте\n", gpu_info)
    else:
        if wmi_gpu := gpu_wmi():
            print_wmic("Информация о видеокарте\n", wmi_gpu)
    if drive_info := hdd_ssd_wmic():
        print_wmic("Информация о HDD и SSD\n", drive_info)
    else:
        if wmi_drive := hdd_ssd_wmi():
            print_wmic("Информация о HDD и SSD\n", wmi_drive)
    if cd_rom_info := cdrom_wmic():
        print_wmic("Информация о CD/DVD-ROM\n", cd_rom_info)
    else:
        if wmi_cdrom := cdrom_wmi():
            print_wmic("Информация о CD/DVD-ROM\n", wmi_cdrom)
    if nic_info := nic_winreg():
        print_wmic("Информация о физических сетевых интерфейсах\n", nic_info)
    else:
        if wmic_nic := nic_wmic():
            print_wmic("Информация о физических сетевых интерфейсах\n", wmic_nic)

    document.add_paragraph(wmic_info)
    document.save(f'{node()}.docx')
    print(f"Собранная информация сохранена в файл: {node()}.docx")
    print(f'\nВремя работы скрипта: {time.monotonic() - t} с.')


if __name__ == "__main__":
    main()
